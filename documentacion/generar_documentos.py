#!/usr/bin/env python3
"""Genera el procedimiento Word y el diagrama Draw.io."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "1F4E79")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading_styled(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=11)
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.5)
    return p


def header_table(table, headers, fill="1F4E79"):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, fill)
        set_cell_border(cell)


def fill_row(row, values, fill=None, bold=False):
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        set_run_font(run, size=10, bold=bold)
        if fill:
            shade_cell(cell, fill)
        set_cell_border(cell)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    header = section.header.paragraphs[0]
    hr = header.add_run("TI-PRO-003  |  Proceso de adquisición de tecnología entre TI y Compras  |  v1.1")
    set_run_font(hr, size=9, color=(0x1F, 0x4E, 0x79))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Documento interno — Área de Tecnología de la Información")
    set_run_font(fr, size=8, color=(0x66, 0x66, 0x66))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROCEDIMIENTO PARA LA GESTIÓN DE ADQUISICIONES TECNOLÓGICAS")
    set_run_font(r, size=18, bold=True, color=(0x1F, 0x4E, 0x79))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Proceso de adquisición de tecnología entre TI y Compras")
    set_run_font(r, size=13, bold=True, color=(0x2E, 0x86, 0xAB))

    meta = doc.add_table(rows=4, cols=4)
    meta.autofit = True
    fill_row(meta.rows[0], ["Código", "TI-PRO-003", "Versión", "1.1"], fill="D6EAF8", bold=True)
    fill_row(meta.rows[1], ["Área responsable", "Tecnología de la Información (TI)", "Áreas involucradas", "TI, Compras, Gerencia"])
    fill_row(meta.rows[2], ["Aplicación", "Clientes internos y externos", "Ciudades", "Quito y Guayaquil"])
    fill_row(
        meta.rows[3],
        ["Equipos típicos", "Laptops, tablets, PCs, accesorios", "Diagrama", "TI-PRO-003-Diagrama.drawio"],
    )

    add_heading_styled(doc, "1. Objetivo", 1)
    add_body(
        doc,
        "Establecer el procedimiento para la adquisición de equipos tecnológicos, accesorios y demás recursos informáticos, "
        "garantizando que todas las compras se realicen de forma planificada, controlada y alineada a las necesidades de la "
        "organización, diferenciando el canal de cliente interno y el canal de cliente externo.",
    )

    add_heading_styled(doc, "2. Alcance", 1)
    add_body(doc, "Aplica a todas las adquisiciones de bienes y servicios tecnológicos requeridos por la organización, incluyendo:")
    for item in [
        "Computadoras y laptops.",
        "Tablets (uso de guardias y supervisores).",
        "Monitores, servidores y equipos de red.",
        "Impresoras y celulares corporativos.",
        "Equipos de videovigilancia y accesorios tecnológicos.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "3. Definición de clientes", 1)
    add_body(
        doc,
        "El proceso tiene dos entradas según quién necesita el equipo. Esta clasificación determina quién solicita la compra a Compras.",
    )

    t = doc.add_table(rows=3, cols=4)
    header_table(t, ["Tipo", "Quiénes son", "Ubicación / canal", "Equipos habituales"])
    fill_row(
        t.rows[1],
        [
            "Cliente interno",
            "Personal administrativo",
            "Quito y Guayaquil (oficinas)",
            "Laptops, PCs, monitores, accesorios de oficina",
        ],
        fill="EAF2F8",
    )
    fill_row(
        t.rows[2],
        [
            "Cliente externo",
            "Todos los clientes por fuera: guardias y supervisores",
            "Operación en campo / sitios externos",
            "Tablets y laptops de operación",
        ],
        fill="E8F8F5",
    )

    add_body(
        doc,
        "Regla de daño o reemplazo: cuando un dispositivo de cliente externo se daña o deja de funcionar, la compra sigue el flujo de cliente externo (solicitud directa a Compras con informe técnico de TI).",
        bold=True,
    )

    add_heading_styled(doc, "4. Responsabilidades", 1)

    add_heading_styled(doc, "4.1 Área solicitante / usuario", 2)
    for item in [
        "Identificar la necesidad o reportar la falla.",
        "Generar el ticket (ambos tipos de cliente).",
        "Cliente interno: esperar la evaluación de TI; no comprar por su cuenta.",
        "Cliente externo: una vez exista el informe técnico de TI, solicitar la compra directamente a Compras por correo.",
        "Firmar el acta de entrega y cuidar el activo asignado.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "4.2 Área de Tecnología de la Información", 2)
    for item in [
        "Analizar la necesidad y evaluar si el equipo tiene arreglo o si existe inventario reutilizable.",
        "Definir especificaciones técnicas y emitir el informe técnico.",
        "Cliente interno: solicitar la compra a Compras cuando no haya arreglo ni stock.",
        "Cliente externo: entregar el informe técnico para que el solicitante gestione la compra con Compras.",
        "Recibir o coordinar el equipo, configurarlo, etiquetarlo e inventariarlo.",
        "Entregar el equipo al usuario y cerrar el ticket.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "4.3 Área de Compras", 2)
    for item in [
        "Recibir la solicitud de compra (de TI en flujo interno, o del cliente/área por correo en flujo externo).",
        "Cotizar, negociar y emitir la orden de compra una vez exista aprobación de Gerencia.",
        "Cliente interno: entregar el equipo a TI.",
        "Cliente externo: entregar el equipo directamente al usuario o notificar por correo para que TI lo prepare.",
        "Dar seguimiento al proveedor y coordinar la entrega.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "4.4 Gerencia", 2)
    add_body(
        doc,
        "Revisa y aprueba la adquisición en ambos flujos (interno y externo), conforme a la política de compras de la empresa. Sin esta aprobación, Compras no emite la orden de compra.",
    )

    add_heading_styled(doc, "5. Pasos a seguir — flujo general", 1)
    add_body(doc, "Estos pasos aplican a toda compra tecnológica. El detalle por tipo de cliente está en las secciones 6 y 7.")

    pasos = [
        ("1. Identificación de la necesidad", "Área solicitante / usuario", "Se detecta necesidad por nuevo colaborador, equipo dañado, obsoleto, incremento de personal o nuevo proyecto. Ejemplo: Tito indica que no le funciona el dispositivo y abre un ticket."),
        ("2. Generación del requerimiento (ticket)", "Área solicitante", "Se registra un ticket o solicitud formal con justificación, usuario final, área, fecha requerida y tipo de recurso."),
        ("3. Evaluación técnica", "TI", "TI analiza inventario, posibilidad de reparación o reutilización, compatibilidad, especificaciones y licencias. Si hay stock o arreglo, no se compra: se repara o reasigna. Si no hay arreglo, se continúa a compra."),
        ("4. Informe y especificaciones técnicas", "TI", "TI elabora el documento técnico (procesador, RAM, SSD, sistema operativo, garantía, accesorios, marca recomendada cuando aplique). Este informe es obligatorio para que Compras proceda."),
        ("5. Solicitud de compra", "TI o cliente externo", "Cliente interno: TI solicita la compra. Cliente externo: el solicitante envía correo a Compras basándose en el informe técnico de TI (debe solicitar directamente)."),
        ("6. Aprobación de Gerencia", "Gerencia", "En los dos casos Gerencia revisa para la aprobación. Si se rechaza, se notifica y se cierra o se reformula el requerimiento."),
        ("7. Orden de compra", "Compras", "Con la aprobación, Compras cotiza (se recomienda al menos tres cotizaciones cuando la política lo requiera), negocia y emite la orden de compra."),
        ("8. Recepción y entrega logística", "Compras", "Cliente interno: Compras entrega el equipo a TI. Cliente externo: Compras entrega el equipo directamente al usuario o envía correo de notificación."),
        ("9. Preparación del equipo", "TI", "En ambos casos TI prepara el equipo: sistema operativo, Active Directory, correo, Microsoft 365, software autorizado, antivirus, políticas de seguridad, inventario y etiquetado."),
        ("10. Entrega al usuario y cierre", "TI y usuario", "El usuario firma el acta de entrega. TI actualiza inventario, licencias y cierra el ticket."),
    ]
    pt = doc.add_table(rows=1 + len(pasos), cols=3)
    header_table(pt, ["Paso", "Responsable", "Qué se hace"])
    for i, (paso, resp, desc) in enumerate(pasos, start=1):
        fill_row(pt.rows[i], [paso, resp, desc], fill="F8FBFD" if i % 2 == 0 else None)

    add_heading_styled(doc, "6. Flujo A — Cliente interno (personal administrativo Quito y Guayaquil)", 1)
    add_body(doc, "Usar este flujo cuando el solicitante es personal administrativo de oficinas en Quito o Guayaquil.")
    for item in [
        "El usuario o su jefatura genera un ticket describiendo la necesidad o la falla.",
        "TI evalúa el caso. Siempre que el equipo tenga arreglo o exista un equipo en inventario, se repara o se reasigna. No se compra.",
        "Solo si no tiene arreglo y no hay stock, TI elabora el informe técnico y solicita la compra a Compras.",
        "Gerencia revisa y aprueba.",
        "Compras procede con la orden de compra y entrega el equipo a TI.",
        "TI prepara el equipo y lo entrega al usuario con acta de entrega.",
        "TI cierra el ticket y actualiza el inventario.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "7. Flujo B — Cliente externo (guardias, supervisores y clientes por fuera)", 1)
    add_body(
        doc,
        "Usar este flujo para guardias, supervisores y demás clientes externos (tablets y laptops de operación). "
        "Si se daña un dispositivo de este grupo, la compra va por este lado.",
    )
    for item in [
        "El usuario reporta que el dispositivo no funciona mediante ticket (ejemplo: Tito indica que no le funciona el dispositivo).",
        "TI evalúa y emite el informe técnico. Si el equipo tiene arreglo, se repara y no se compra.",
        "Si corresponde compra, el cliente o su área debe solicitar directamente a Compras, enviando un correo basado en el informe técnico de TI. No espera a que TI haga el pedido en su nombre.",
        "Gerencia revisa y aprueba (igual que en el flujo interno).",
        "Compras procede con la orden de compra.",
        "Compras entrega el equipo directamente al usuario o envía un correo de notificación / coordinación.",
        "TI prepara el equipo (configuración, seguridad e inventario) para dejarlo listo.",
        "Se entrega o habilita el equipo al usuario y se cierra el ticket.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "8. Criterio de decisión: ¿se compra o no?", 1)
    dt = doc.add_table(rows=4, cols=2)
    header_table(dt, ["Situación", "Acción"])
    fill_row(dt.rows[1], ["El equipo tiene arreglo o hay inventario disponible", "No se compra. TI repara o reasigna y entrega."])
    fill_row(dt.rows[2], ["No tiene arreglo y no hay stock — cliente interno", "TI solicita la compra. Gerencia aprueba. Compras entrega a TI. TI prepara y entrega al usuario."], fill="EAF2F8")
    fill_row(dt.rows[3], ["No tiene arreglo y no hay stock — cliente externo / equipo dañado de campo", "TI emite informe. El solicitante escribe a Compras. Gerencia aprueba. Compras entrega al usuario o notifica por correo. TI prepara el equipo."], fill="E8F8F5")

    add_heading_styled(doc, "9. Preparación del equipo (TI) — actividades mínimas", 1)
    for item in [
        "Verificar estado físico, número de serie, modelo, garantía, accesorios y funcionamiento. Si hay inconsistencias, notificar al proveedor.",
        "Instalar sistema operativo, Active Directory, correo corporativo y Microsoft 365.",
        "Instalar software autorizado y antivirus; aplicar políticas de seguridad.",
        "Registrar en inventario y etiquetar el activo.",
        "Entregar con acta: equipo, accesorios, estado, fecha, responsable de TI y usuario receptor.",
        "Actualizar inventario, base de equipos, licencias y cerrar el ticket.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "10. Documentos del proceso", 1)
    docs = [
        ("Ticket / solicitud", "Inicio del requerimiento"),
        ("Informe técnico de TI", "Base obligatoria para cotizar y comprar"),
        ("Correo de solicitud a Compras", "Flujo externo: lo envía el cliente/área; flujo interno: lo gestiona TI"),
        ("Aprobación de Gerencia", "Autorización para emitir la orden de compra"),
        ("Orden de compra", "Documento de Compras hacia el proveedor"),
        ("Acta de entrega del equipo", "Cierre con el usuario"),
    ]
    d2 = doc.add_table(rows=1 + len(docs), cols=2)
    header_table(d2, ["Documento", "Para qué sirve"])
    for i, (n, u) in enumerate(docs, start=1):
        fill_row(d2.rows[i], [n, u])

    add_heading_styled(doc, "11. Diagrama", 1)
    add_body(
        doc,
        "El diagrama de flujo para Draw.io se encuentra en el archivo TI-PRO-003-Diagrama.drawio. "
        "Ábralo en https://app.diagrams.net o en la aplicación de escritorio Draw.io / diagrams.net.",
    )

    add_heading_styled(doc, "12. Control de cambios", 1)
    cc = doc.add_table(rows=3, cols=4)
    header_table(cc, ["Versión", "Fecha", "Cambio", "Autor"])
    fill_row(cc.rows[1], ["1.0", "—", "Procedimiento general de adquisiciones tecnológicas", "Área de TI"])
    fill_row(cc.rows[2], ["1.1", "2026-08-28", "Se incorporan flujos de cliente interno (Quito/GYE) y cliente externo (guardias/supervisores), regla de daño, solicitud directa a Compras y entrega.", "Área de TI"])

    path = OUT / "TI-PRO-003-Proceso-Adquisicion-Tecnologia.docx"
    doc.save(path)
    return path


def mx_cell(cid, value, style, x, y, w, h, parent="1", vertex="1"):
    value_esc = (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("\n", "&#xa;")
    )
    # Allow simple HTML if already escaped via using raw html markers
    return (
        f'        <mxCell id="{cid}" value="{value_esc}" style="{style}" vertex="{vertex}" parent="{parent}">\n'
        f'          <mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/>\n'
        f"        </mxCell>\n"
    )


def mx_html(cid, html, style, x, y, w, h, parent="1"):
    return (
        f'        <mxCell id="{cid}" value="{html}" style="{style}" vertex="1" parent="{parent}">\n'
        f'          <mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/>\n'
        f"        </mxCell>\n"
    )


def mx_edge(eid, source, target, style="edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;strokeWidth=2;strokeColor=#34495E;endArrow=block;endFill=1;", label=""):
    lab = f' value="{label}"' if label else ' value=""'
    return (
        f'        <mxCell id="{eid}"{lab} style="{style}" edge="1" parent="1" source="{source}" target="{target}">\n'
        f'          <mxGeometry relative="1" as="geometry"/>\n'
        f"        </mxCell>\n"
    )


def build_drawio():
    # Styles
    title_s = "rounded=0;whiteSpace=wrap;html=1;fillColor=#1F4E79;fontColor=#FFFFFF;fontStyle=1;fontSize=16;strokeColor=#1F4E79;"
    subtitle_s = "rounded=0;whiteSpace=wrap;html=1;fillColor=#D6EAF8;fontColor=#1F4E79;fontSize=12;strokeColor=#5DADE2;"
    lane_i = "rounded=0;whiteSpace=wrap;html=1;fillColor=#EAF2F8;strokeColor=#1F4E79;strokeWidth=2;fontStyle=1;fontSize=14;fontColor=#1F4E79;verticalAlign=top;spacingTop=8;"
    lane_e = "rounded=0;whiteSpace=wrap;html=1;fillColor=#E8F8F5;strokeColor=#117A65;strokeWidth=2;fontStyle=1;fontSize=14;fontColor=#0E6655;verticalAlign=top;spacingTop=8;"
    start_s = "ellipse;whiteSpace=wrap;html=1;fillColor=#1ABC9C;fontColor=#FFFFFF;strokeColor=#0E6655;fontStyle=1;fontSize=13;strokeWidth=2;"
    end_s = "ellipse;whiteSpace=wrap;html=1;fillColor=#922B21;fontColor=#FFFFFF;strokeColor=#641E16;fontStyle=1;fontSize=13;strokeWidth=2;"
    dec_s = "rhombus;whiteSpace=wrap;html=1;fillColor=#F9E79F;strokeColor=#B7950B;fontColor=#7D6608;fontStyle=1;fontSize=12;strokeWidth=2;"
    ti_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#D7BDE2;strokeColor=#6C3483;fontColor=#4A235A;fontStyle=0;fontSize=11;strokeWidth=2;arcSize=12;"
    int_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#AED6F1;strokeColor=#1F4E79;fontColor=#1B4F72;fontSize=11;strokeWidth=2;arcSize=12;"
    ext_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#A3E4D7;strokeColor=#0E6655;fontColor=#0E6655;fontSize=11;strokeWidth=2;arcSize=12;"
    ger_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#F9E79F;strokeColor=#B7950B;fontColor=#7D6608;fontSize=11;strokeWidth=2;fontStyle=1;arcSize=12;"
    com_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#F5B7B1;strokeColor=#922B21;fontColor=#641E16;fontSize=11;strokeWidth=2;arcSize=12;"
    ok_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#ABEBC6;strokeColor=#1D8348;fontColor=#145A32;fontSize=11;strokeWidth=2;arcSize=12;"
    no_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#F5B7B1;strokeColor=#922B21;fontColor=#641E16;fontSize=11;strokeWidth=2;arcSize=12;"
    note_s = "shape=note;whiteSpace=wrap;html=1;size=18;fillColor=#FCF3CF;strokeColor=#B7950B;fontSize=10;align=left;spacingLeft=8;fontColor=#7D6608;"
    legend_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#85929E;fontSize=10;align=left;spacingLeft=10;arcSize=8;"
    merge_s = "rounded=1;whiteSpace=wrap;html=1;fillColor=#FAD7A0;strokeColor=#B9770E;fontColor=#6E2C00;fontStyle=1;fontSize=12;strokeWidth=2;arcSize=10;"

    yes_e = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;strokeWidth=2;strokeColor=#1D8348;endArrow=block;endFill=1;fontColor=#1D8348;fontStyle=1;fontSize=11;"
    no_e = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;strokeWidth=2;strokeColor=#922B21;endArrow=block;endFill=1;fontColor=#922B21;fontStyle=1;fontSize=11;"
    def_e = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;strokeWidth=2;strokeColor=#34495E;endArrow=block;endFill=1;"
    blue_e = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;strokeWidth=2;strokeColor=#1F4E79;endArrow=block;endFill=1;fontColor=#1F4E79;fontStyle=1;"
    green_e = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;strokeWidth=2;strokeColor=#0E6655;endArrow=block;endFill=1;fontColor=#0E6655;fontStyle=1;"

    cells = []
    cells.append('        <mxCell id="0"/>\n')
    cells.append('        <mxCell id="1" parent="0"/>\n')

    # Title
    cells.append(mx_cell("t1", "PROCESO DE ADQUISICIÓN DE TECNOLOGÍA ENTRE TI Y COMPRAS", title_s, 40, 20, 1520, 50))
    cells.append(mx_cell("t2", "TI-PRO-003  ·  v1.1  ·  Clientes internos (admin. Quito y GYE) y externos (guardias / supervisores — tablets y laptops)", subtitle_s, 40, 70, 1520, 36))

    # Legend
    cells.append(mx_cell("lg", "Leyenda de responsables:  Azul = Cliente interno    Verde = Cliente externo    Lila = TI    Amarillo = Gerencia    Rojo = Compras    Rombo = decisión", legend_s, 40, 116, 1520, 36))

    # Start
    cells.append(mx_cell("start", "INICIO\nNecesidad o falla\nde un equipo", start_s, 680, 170, 200, 80))

    cells.append(mx_cell("tipo", "¿Quién solicita?\n¿Cliente interno o\ncliente externo?", dec_s, 650, 290, 260, 130))

    # Lanes
    cells.append(mx_cell("laneI", "FLUJO A — CLIENTE INTERNO\nPersonal administrativo · Quito y Guayaquil", lane_i, 40, 460, 720, 720))
    cells.append(mx_cell("laneE", "FLUJO B — CLIENTE EXTERNO\nGuardias, supervisores y clientes por fuera · Tablets / laptops", lane_e, 840, 460, 720, 720))

    # Internal flow
    cells.append(mx_cell("i1", "1. Genera TICKET\nde necesidad o falla", int_s, 80, 530, 280, 70))
    cells.append(mx_cell("i2", "2. TI evalúa el caso\n(inventario, compatibilidad, arreglo)", ti_s, 80, 630, 280, 70))
    cells.append(mx_cell("i3", "¿El equipo tiene arreglo\no hay inventario?", dec_s, 90, 730, 260, 110))
    cells.append(mx_cell("i4", "TI repara o reasigna.\nNo se compra.", ok_s, 430, 750, 280, 70))
    cells.append(mx_cell("i5", "3. TI elabora informe técnico\ny SOLICITA la compra a Compras", ti_s, 80, 880, 300, 80))
    cells.append(mx_cell("i6", "Entrega al usuario y cierra ticket\n(sin compra)", ok_s, 430, 860, 280, 70))
    cells.append(mx_cell("inote", "Siempre que el equipo tenga arreglo, no se compra. TI pide la compra solo si no hay solución técnica ni stock.", note_s, 80, 990, 640, 70))
    cells.append(mx_cell("i7", "Continúa a aprobación\nde Gerencia →", merge_s, 80, 1080, 280, 60))

    # External flow
    cells.append(mx_cell("e1", "1. Reporta con TICKET\nEj.: Tito — no le funciona el dispositivo", ext_s, 880, 530, 340, 70))
    cells.append(mx_cell("e2", "2. TI evalúa y emite\nINFORME TÉCNICO", ti_s, 880, 630, 280, 70))
    cells.append(mx_cell("e3", "¿El equipo tiene arreglo\no hay inventario?", dec_s, 890, 730, 260, 110))
    cells.append(mx_cell("e4", "TI repara o reasigna.\nNo se compra.", ok_s, 1220, 750, 300, 70))
    cells.append(mx_cell("e5", "3. El cliente SOLICITA DIRECTAMENTE\na Compras por CORREO, con el\ninforme técnico de TI", ext_s, 880, 870, 340, 90))
    cells.append(mx_cell("e6", "Entrega al usuario y cierra ticket\n(sin compra)", ok_s, 1220, 860, 300, 70))
    cells.append(mx_cell("enote", "Si se daña un equipo de campo (guardia/supervisor), la compra va por este flujo. El usuario no espera a que TI haga el pedido.", note_s, 880, 990, 640, 70))
    cells.append(mx_cell("e7", "Continúa a aprobación\nde Gerencia →", merge_s, 880, 1080, 280, 60))

    # Shared gerencia / compras
    cells.append(mx_cell("g1", "4. GERENCIA revisa para aprobación\n(aplica en los DOS casos)", ger_s, 560, 1220, 480, 70))
    cells.append(mx_cell("g2", "¿Gerencia aprueba\nla compra?", dec_s, 650, 1320, 260, 110))
    cells.append(mx_cell("gno", "Se notifica el rechazo.\nSe reformula o se cierra\nel requerimiento.", no_s, 1000, 1335, 280, 80))
    cells.append(mx_cell("c1", "5. COMPRAS procede con cotización\ny emite la ORDEN DE COMPRA", com_s, 560, 1470, 480, 70))

    cells.append(mx_cell("ent", "6. ¿Cómo se entrega el equipo\nrecibido del proveedor?", dec_s, 640, 1570, 280, 120))

    # Delivery split
    cells.append(mx_cell("dI1", "Cliente interno:\nCompras ENTREGA el equipo a TI", com_s, 80, 1740, 320, 70))
    cells.append(mx_cell("dI2", "7. TI PREPARA el equipo\n(SO, AD, correo, M365, antivirus,\npolíticas, inventario, etiqueta)", ti_s, 80, 1840, 320, 90))
    cells.append(mx_cell("dI3", "8. TI entrega al usuario\ncon Acta de entrega", int_s, 80, 1960, 320, 70))

    cells.append(mx_cell("dE1", "Cliente externo:\nCompras entrega DIRECTAMENTE al usuario\no ENVÍA CORREO de coordinación", com_s, 1160, 1740, 380, 80))
    cells.append(mx_cell("dE2", "7. TI PREPARA el equipo\npara dejarlo listo de uso", ti_s, 1190, 1850, 320, 80))
    cells.append(mx_cell("dE3", "8. Se habilita / entrega al usuario\ncon Acta de entrega", ext_s, 1190, 1960, 320, 70))

    cells.append(mx_cell("fin1", "9. CIERRE\nInventario · licencias · ticket", ok_s, 620, 2080, 320, 70))
    cells.append(mx_cell("end", "FIN", end_s, 700, 2180, 160, 60))

    # Edges
    cells.append(mx_edge("a1", "start", "tipo", def_e))
    cells.append(mx_edge("a2", "tipo", "i1", blue_e, "Interno"))
    cells.append(mx_edge("a3", "tipo", "e1", green_e, "Externo"))

    cells.append(mx_edge("a4", "i1", "i2", def_e))
    cells.append(mx_edge("a5", "i2", "i3", def_e))
    cells.append(mx_edge("a6", "i3", "i4", yes_e, "Sí"))
    cells.append(mx_edge("a7", "i3", "i5", no_e, "No"))
    cells.append(mx_edge("a8", "i4", "i6", def_e))
    cells.append(mx_edge("a9", "i5", "i7", def_e))
    cells.append(mx_edge("a10", "i6", "fin1", yes_e))

    cells.append(mx_edge("b1", "e1", "e2", def_e))
    cells.append(mx_edge("b2", "e2", "e3", def_e))
    cells.append(mx_edge("b3", "e3", "e4", yes_e, "Sí"))
    cells.append(mx_edge("b4", "e3", "e5", no_e, "No"))
    cells.append(mx_edge("b5", "e4", "e6", def_e))
    cells.append(mx_edge("b6", "e5", "e7", def_e))
    cells.append(mx_edge("b7", "e6", "fin1", yes_e))

    cells.append(mx_edge("c2", "i7", "g1", def_e))
    cells.append(mx_edge("c3", "e7", "g1", def_e))
    cells.append(mx_edge("c4", "g1", "g2", def_e))
    cells.append(mx_edge("c5", "g2", "gno", no_e, "No"))
    cells.append(mx_edge("c6", "g2", "c1", yes_e, "Sí"))
    cells.append(mx_edge("c7", "gno", "end", no_e))
    cells.append(mx_edge("c8", "c1", "ent", def_e))

    cells.append(mx_edge("d1", "ent", "dI1", blue_e, "Interno"))
    cells.append(mx_edge("d2", "ent", "dE1", green_e, "Externo"))
    cells.append(mx_edge("d3", "dI1", "dI2", def_e))
    cells.append(mx_edge("d4", "dI2", "dI3", def_e))
    cells.append(mx_edge("d5", "dI3", "fin1", def_e))
    cells.append(mx_edge("d6", "dE1", "dE2", def_e))
    cells.append(mx_edge("d7", "dE2", "dE3", def_e))
    cells.append(mx_edge("d8", "dE3", "fin1", def_e))
    cells.append(mx_edge("d9", "fin1", "end", def_e))

    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<mxfile host="app.diagrams.net" agent="Cursor" version="22.1.0" type="device">\n'
        '  <diagram id="ti-pro-003" name="Adquisición TI y Compras">\n'
        '    <mxGraphModel dx="1400" dy="900" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" '
        'arrows="1" fold="1" page="1" pageScale="1" pageWidth="1640" pageHeight="2300" math="0" shadow="0">\n'
        "      <root>\n"
        + "".join(cells)
        + "      </root>\n"
        "    </mxGraphModel>\n"
        "  </diagram>\n"
        "</mxfile>\n"
    )
    path = OUT / "TI-PRO-003-Diagrama.drawio"
    path.write_text(xml, encoding="utf-8")
    return path


def build_markdown():
    md = """# Procedimiento para la gestión de adquisiciones tecnológicas

**Código:** TI-PRO-003  
**Versión:** 1.1  
**Proceso:** Adquisición de tecnología entre TI y Compras  
**Diagrama Draw.io:** `TI-PRO-003-Diagrama.drawio` (abrir en [diagrams.net](https://app.diagrams.net))

## 1. Objetivo

Establecer el procedimiento para adquirir equipos tecnológicos de forma planificada y controlada, diferenciando el canal de **cliente interno** y el de **cliente externo**.

## 2. Quién es cada cliente

| Tipo | Quiénes son | Dónde | Equipos habituales |
| --- | --- | --- | --- |
| **Cliente interno** | Personal administrativo | Quito y Guayaquil | Laptops, PCs, monitores, accesorios |
| **Cliente externo** | Todos los clientes por fuera: guardias y supervisores | Operación / campo | Tablets y laptops |

**Regla de daño:** si se daña un dispositivo de cliente externo, la compra va por el flujo externo.

## 3. Pasos a seguir (resumen)

1. El usuario identifica la necesidad o la falla (ejemplo: Tito indica que no le funciona el dispositivo).
2. Se genera un **ticket**.
3. **TI evalúa.** Si tiene arreglo o hay inventario, se repara o se reasigna. **No se compra.**
4. Si no hay arreglo, TI emite el **informe técnico**.
5. **Solicitud de compra:**
   - Interno: **TI solicita** la compra a Compras.
   - Externo: el cliente **solicita directamente** a Compras por **correo**, con el informe de TI.
6. **Gerencia revisa y aprueba en los dos casos.**
7. Compras cotiza y emite la **orden de compra**.
8. Entrega logística:
   - Interno: Compras **entrega el equipo a TI**.
   - Externo: Compras **entrega al usuario** o **envía correo**.
9. **TI prepara el equipo** (configuración, seguridad, inventario).
10. Entrega al usuario con **acta**, actualización de inventario y **cierre del ticket**.

## 4. Flujo A — Cliente interno (admin. Quito y GYE)

1. Ticket de necesidad o falla.
2. Evaluación por TI. Siempre que no tenga arreglo (y no haya stock) se pasa a compra.
3. TI solicita la compra con informe técnico.
4. Gerencia aprueba.
5. Compras emite la orden y entrega a TI.
6. TI prepara el equipo y lo entrega al usuario.

## 5. Flujo B — Cliente externo (guardias / supervisores)

1. Ticket (dispositivo no funciona).
2. TI evalúa y emite informe técnico.
3. El área/usuario envía correo a Compras basándose en el informe. Debe solicitar directamente.
4. Gerencia aprueba.
5. Compras procede con la orden.
6. Compras entrega el equipo directamente al usuario o envía correo.
7. TI prepara el equipo para entregarlo / habilitarlo.

## 6. Cómo usar el diagrama

1. Abra [https://app.diagrams.net](https://app.diagrams.net).
2. Elija **Abrir archivo existente** y seleccione `TI-PRO-003-Diagrama.drawio`.
3. Edite textos, responsables o tiempos según la política interna.
4. Exporte a PNG o PDF desde **Archivo → Exportar**.
"""
    path = OUT / "TI-PRO-003-Pasos-a-seguir.md"
    path.write_text(md, encoding="utf-8")
    return path


if __name__ == "__main__":
    from aplicar_formato_plantilla import build as build_docx_plantilla

    p1 = build_docx_plantilla()
    p2 = build_drawio()
    p3 = build_markdown()
    print(p1)
    print(p2)
    print(p3)
